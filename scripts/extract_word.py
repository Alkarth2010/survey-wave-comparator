from docx import Document
import pandas as pd
import re

def extract_questions(docx_path, wave):
    doc = Document(docx_path)
    rows = []

    # -------------------------------------------------
    # EXISTING LOGIC (UNCHANGED)
    # -------------------------------------------------
    for p in doc.paragraphs:
        txt = p.text.strip()
        if len(txt) > 10 and "?" in txt:
            rows.append({
                "question_text": txt,
                "wave": wave
            })

    # -------------------------------------------------
    # NEW ADDITION: HANDLE TABLE-BASED QUESTIONNAIRES
    # -------------------------------------------------
    qcode_pattern = re.compile(r"^[A-Z]\d+[a-z]?$", re.I)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]

            if len(cells) < 2:
                continue

            qcode = cells[0]
            qtext = cells[1]

            # Detect question rows like S1e, S1f, Q12
            if (
                qcode_pattern.match(qcode)
                and len(qtext) > 20
                and not qtext.lower().startswith(("yes", "no", "once", "less"))
            ):
                rows.append({
                    "question_text": qtext,
                    "wave": wave
                })

    return pd.DataFrame(rows, columns=["question_text", "wave"])
